from fpdf import FPDF
from docx import Document

def export_to_pdf(sections, filepath):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for title, content, tags in sections:
        pdf.set_font("Arial", style='B', size=14)
        pdf.cell(0, 10, title, ln=1)
        pdf.set_font("Arial", size=12)
        for line in content.split('\n'):
            pdf.multi_cell(0, 10, line)
        pdf.ln(5)
    pdf.output(filepath)

def export_to_docx(sections, filepath):
    doc = Document()
    for title, content, tags in sections:
        doc.add_heading(title, level=1)
        for line in content.split('\n'):
            run = doc.add_paragraph().add_run(line)
            # Basic formatting from tags (expand as needed)
            if tags.get('bold'):
                run.bold = True
            if tags.get('italic'):
                run.italic = True
            if tags.get('underline'):
                run.underline = True
    doc.save(filepath)